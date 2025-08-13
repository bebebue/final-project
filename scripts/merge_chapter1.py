import os
import re
import shutil
import sys
from difflib import SequenceMatcher
from typing import List, Dict, Optional

try:
	from docx import Document
	from docx.text.paragraph import Paragraph
	from docx.document import Document as _Document
	from docx.shared import Pt
except Exception as e:
	print("Missing dependency: python-docx. Install with 'pip install python-docx'", file=sys.stderr)
	raise


def normalize_title(text: str) -> str:
	if text is None:
		return ""
	# Lowercase, strip, remove numbers (Arabic and Thai), most punctuation, and extra spaces
	text = text.strip()
	text = text.lower()
	# Remove Arabic digits and dots like 1., 1.1, 1.1.1
	text = re.sub(r"\b\d+(?:[\.-]\d+)*\.?\s*", "", text)
	# Remove Thai digits
	text = re.sub(r"[\u0E50-\u0E59]", "", text)
	# Remove common punctuation
	text = re.sub(r"[\-–—•·\(\)\[\]\{\}:;,.]", " ", text)
	# Collapse whitespace
	text = re.sub(r"\s+", " ", text)
	return text.strip()


def _heading_level_from_text(text: str) -> Optional[int]:
	if not text:
		return None
	raw = text.strip()
	# Normalize Thai digits to Arabic 0-9 for counting dots
	thai_to_arabic = str.maketrans("๐๑๒๓๔๕๖๗๘๙", "0123456789")
	raw_norm = raw.translate(thai_to_arabic)
	# Match patterns like "บทที่ 1", "1.", "1.1", "1.1.1" etc
	if re.match(r"^(บทที่)\s*\d+", raw_norm):
		return 1
	m = re.match(r"^(\d+)(?:\.(\d+))*[\.)\s]", raw_norm)
	if m:
		# Level equals number of numeric components
		parts = re.findall(r"\d+", raw_norm.split()[0])
		if parts:
			return max(1, len(parts))
	return None


def get_heading_level(paragraph: Paragraph) -> Optional[int]:
	style_name = paragraph.style.name if paragraph.style is not None else ""
	# Prefer style-based detection
	m = re.search(r"(\d+)$", style_name)
	if style_name.startswith("Heading") and m:
		try:
			return int(m.group(1))
		except ValueError:
			return None
	# Localized heading styles containing digits
	if "หัวเรื่อง" in style_name:
		m2 = re.search(r"(\d+)", style_name)
		if m2:
			try:
				return int(m2.group(1))
			except ValueError:
				return None
	# Fallback to text-based heuristic
	text_level = _heading_level_from_text(paragraph.text)
	if text_level is not None:
		return text_level
	return None


def parse_sections(doc: _Document) -> List[Dict]:
	sections: List[Dict] = []
	current: Optional[Dict] = None
	for p in doc.paragraphs:
		level = get_heading_level(p)
		if level is not None:
			# Start a new section
			current = {
				"title": p.text.strip(),
				"norm_title": normalize_title(p.text),
				"level": level,
				"paragraphs": []
			}
			sections.append(current)
		else:
			if current is not None:
				current["paragraphs"].append(p)
	return sections


def build_title_index(sections: List[Dict]) -> Dict[str, int]:
	index: Dict[str, int] = {}
	for i, sec in enumerate(sections):
		key = sec.get("norm_title", "")
		if key:
			index[key] = i
	return index


def find_best_match(template_title: str, source_index: Dict[str, int]) -> Optional[int]:
	t_key = normalize_title(template_title)
	if t_key in source_index:
		return source_index[t_key]
	# Fuzzy match
	best_idx = None
	best_score = 0.0
	for s_key, s_idx in source_index.items():
		score = SequenceMatcher(a=t_key, b=s_key).ratio()
		if score > best_score:
			best_score = score
			best_idx = s_idx
	# Accept only reasonably similar
	if best_score >= 0.6:
		return best_idx
	return None


def copy_run_formatting(src_run, dst_run):
	try:
		dst_run.bold = src_run.bold
		dst_run.italic = src_run.italic
		dst_run.underline = src_run.underline
		if src_run.font is not None:
			dst_run.font.name = src_run.font.name
			dst_run.font.size = src_run.font.size
	except Exception:
		pass


def add_paragraph_with_style(dst_doc: _Document, text: str = "", style_name: Optional[str] = None):
	p = dst_doc.add_paragraph()
	if style_name:
		try:
			p.style = dst_doc.styles[style_name]
		except Exception:
			# Fallback silently if style not found
			pass
	if text:
		run = p.add_run(text)
	return p


def merge_chapter(template_path: str, source_path: str, target_path: str) -> Dict[str, int]:
	if not os.path.exists(template_path):
		raise FileNotFoundError(f"Template not found: {template_path}")
	if not os.path.exists(source_path):
		raise FileNotFoundError(f"Source not found: {source_path}")

	template_doc = Document(template_path)
	source_doc = Document(source_path)

	template_sections = parse_sections(template_doc)
	source_sections = parse_sections(source_doc)
	source_index = build_title_index(source_sections)

	out_doc = Document()

	matched = 0
	skipped = 0

	for t_sec in template_sections:
		# Add heading
		heading_text = t_sec["title"]
		level = t_sec["level"]
		# Use built-in heading style
		try:
			out_doc.add_heading(heading_text, level=level)
		except Exception:
			# Fallback to style name
			style_name = f"Heading {level}"
			add_paragraph_with_style(out_doc, heading_text, style_name)

		# Find matching source section
		s_idx = find_best_match(heading_text, source_index)
		if s_idx is None:
			# keep empty section
			skipped += 1
			continue

		s_sec = source_sections[s_idx]
		# Copy paragraphs
		for sp in s_sec["paragraphs"]:
			np = out_doc.add_paragraph()
			# Try to copy style name
			try:
				style_name = sp.style.name if sp.style is not None else None
				if style_name:
					np.style = out_doc.styles[style_name]
			except Exception:
				pass
			# Copy runs with basic formatting
			if sp.runs:
				for r in sp.runs:
					new_run = np.add_run(r.text)
					copy_run_formatting(r, new_run)
			else:
				# Plain text paragraph
				np.add_run(sp.text)

		matched += 1

	# Backup and write target
	if os.path.exists(target_path):
		base, ext = os.path.splitext(target_path)
		backup_path = f"{base} (backup){ext}"
		try:
			shutil.copy2(target_path, backup_path)
			print(f"Backed up existing target to: {backup_path}")
		except Exception as e:
			print(f"Warning: failed to backup target: {e}", file=sys.stderr)

	out_doc.save(target_path)
	print(f"Wrote merged Chapter 1 to: {target_path}")

	return {"matched_sections": matched, "skipped_sections": skipped, "total_template_sections": len(template_sections)}


def main():
	# Default paths for this workspace
	template_path = "/workspace/Full paper/5. บทที่ 1.docx"
	source_path = "/workspace/ref/ข้อเสนอโครงงาน.docx"
	target_path = "/workspace/5.บทที่ 1.docx"

	# Allow overrides via CLI args
	if len(sys.argv) >= 2:
		template_path = sys.argv[1]
	if len(sys.argv) >= 3:
		source_path = sys.argv[2]
	if len(sys.argv) >= 4:
		target_path = sys.argv[3]

	stats = merge_chapter(template_path, source_path, target_path)
	print("Stats:", stats)


if __name__ == "__main__":
	main()